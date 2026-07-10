"""
知识入库器（Indexer）

把竞赛规则数据源（docx 表格、xlsx 等）解析、切分、向量化后写入 Chroma。

切分策略
========
竞赛等级分类表是表格结构（序号 | 竞赛名称 | 级别 | 类别 | 备注），
每个竞赛天然是一个独立的"知识条目"，因此采用【逐行切分】而非递归字符切分：
- 每行 → 一个 Document，page_content 是自然语言描述（便于语义检索）
- metadata 含 竞赛名称/级别/类别/序号（便于精确过滤）

这样的好处：
1. 检索粒度精确（问"挑战杯是什么级别"直接命中该行）
2. 元数据过滤（可按级别/类别筛选）
3. 避免递归切分把多个竞赛切碎导致语义混乱
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def _parse_competition_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    解析竞赛等级分类表 docx，返回知识条目列表。

    docx 结构：1 个表格，列 = [序号, 竞赛名称, 级别, 类别, 备注]
    （前两行是合并的标题行，需跳过）

    Returns:
        [{"name", "level", "category", "remark", "seq"}, ...]
    """
    import docx  # python-docx，已在项目依赖中

    doc = docx.Document(docx_path)
    if not doc.tables:
        logger.warning("docx 中无表格: %s", docx_path)
        return []

    table = doc.tables[0]
    entries: List[Dict[str, Any]] = []

    for row_idx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        # 跳过表头（识别"序号"或"竞赛名称"字样的行）
        joined = "".join(cells)
        if row_idx < 3 and ("序号" in joined or "竞赛名称" in joined or "等级" in joined):
            continue
        # 容错：至少需要 4 列有效内容
        if len(cells) < 4:
            continue
        seq, name, level, category = cells[0], cells[1], cells[2], cells[3]
        remark = cells[4] if len(cells) > 4 else ""
        # 跳过空行
        if not name:
            continue
        entries.append({
            "seq": seq,
            "name": name,
            "level": level,
            "category": category,
            "remark": remark,
        })

    logger.info("从 %s 解析出 %d 条竞赛条目", Path(docx_path).name, len(entries))
    return entries


def _entries_to_documents(entries: List[Dict[str, Any]]) -> List[Any]:
    """
    把竞赛条目转为 LangChain Document 列表。

    page_content 用自然语言描述（利于 embedding 语义检索）：
        "竞赛名称：挑战杯全国大学生课外学术科技作品竞赛。级别：国家级。类别：A类赛事。"

    metadata 含结构化字段（利于过滤）。
    """
    from langchain_core.documents import Document

    docs: List[Document] = []
    for e in entries:
        parts = [f"竞赛名称：{e['name']}"]
        if e["level"]:
            parts.append(f"级别：{e['level']}")
        if e["category"]:
            parts.append(f"类别：{e['category']}类赛事")
        if e["remark"]:
            parts.append(f"备注：{e['remark']}")
        content = "。".join(parts) + "。"
        docs.append(Document(
            page_content=content,
            metadata={
                "name": e["name"],
                "level": e["level"],
                "category": e["category"],
                "remark": e["remark"],
                "seq": e["seq"],
                "source": "competition_levels",
            },
        ))
    return docs


def index_competition_levels(
    config_loader,
    vectorstore,
    *,
    docx_path: Optional[str] = None,
    clear_existing: bool = True,
) -> int:
    """
    将竞赛等级分类表入库到向量库。

    Args:
        config_loader: ConfigLoader 实例
        vectorstore: Chroma 向量库实例
        docx_path: docx 路径；None 则从配置 rag.knowledge_sources.competition_levels_doc 读取
        clear_existing: 入库前是否清空该集合中 source=competition_levels 的旧数据（避免重复）

    Returns:
        入库的文档数

    Raises:
        FileNotFoundError: docx 不存在
        ValueError: 配置缺失
    """
    if docx_path is None:
        config = config_loader.load_config()
        rel = config.get("rag", {}).get("knowledge_sources", {}).get("competition_levels_doc")
        if not rel:
            raise ValueError("rag.knowledge_sources.competition_levels_doc 未配置")
        docx_path = str((config_loader.project_root / rel).resolve())

    if not Path(docx_path).exists():
        raise FileNotFoundError(f"竞赛等级分类表不存在: {docx_path}")

    entries = _parse_competition_docx(docx_path)
    if not entries:
        logger.warning("未解析到任何竞赛条目，跳过入库")
        return 0

    docs = _entries_to_documents(entries)

    # 清空旧数据（按 source 过滤，避免重复入库导致检索重复）
    if clear_existing:
        try:
            existing = vectorstore.get(where={"source": "competition_levels"})
            if existing and existing.get("ids"):
                vectorstore.delete(ids=existing["ids"])
                logger.info("已清除旧竞赛条目 %d 条", len(existing["ids"]))
        except Exception as e:
            # 旧数据清除失败不阻断入库（可能是空库）
            logger.debug("清除旧数据时: %s", e)

    # 写入（分批：智谱等 embedding API 单次最多 64 条，留余量取 32/批）
    ids = [f"comp_{e['seq']}" for e in entries]
    _add_documents_batched(vectorstore, docs, ids, batch_size=32)
    logger.info("竞赛等级分类表入库完成：%d 条", len(docs))
    return len(docs)


def _add_documents_batched(vectorstore, docs, ids, batch_size: int = 32):
    """
    分批写入向量库。

    智谱 embedding-3 等部分 Provider 限制单次 embedding 请求最多 64 条输入，
    Chroma 默认批量提交会超出限制，因此按 batch_size 分批 add_documents。
    这是 RAG 工程针对具体 Provider 约束的常见适配。
    """
    total = len(docs)
    for start in range(0, total, batch_size):
        end = start + batch_size
        batch_docs = docs[start:end]
        batch_ids = ids[start:end]
        vectorstore.add_documents(documents=batch_docs, ids=batch_ids)
        logger.debug("已入库 %d/%d", end, total)


def get_collection_stats(vectorstore) -> Dict[str, Any]:
    """返回向量库的简单统计信息。"""
    try:
        data = vectorstore.get()
        return {
            "total": len(data.get("ids", [])),
            "ids_sample": data.get("ids", [])[:5],
        }
    except Exception as e:
        logger.error("获取统计失败: %s", e)
        return {"total": 0, "error": str(e)}


__all__ = [
    "index_competition_levels",
    "get_collection_stats",
    "_add_documents_batched",
]
